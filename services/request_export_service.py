"""
Сервис экспорта заявок в различные форматы.

Предоставляет функциональность для:
- Экспорт в PDF с профессиональным оформлением
- Экспорт в DOCX с настраиваемыми шаблонами
- Экспорт в XLSX с данными для анализа
- Пакетный экспорт множества заявок
- Настраиваемые шаблоны экспорта
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Union

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import black, blue, red, green
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# XLSX generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from utils.logger import get_logger
from utils.exceptions import ValidationError, BusinessLogicError

logger = get_logger(__name__)


class RequestExportService:
    """
    Сервис экспорта заявок лаборатории.
    """
    
    def __init__(self, db_connection, export_dir: str = "exports"):
        """
        Инициализация сервиса.
        
        Args:
            db_connection: Подключение к базе данных
            export_dir: Директория для экспорта файлов
        """
        self.db_connection = db_connection
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(exist_ok=True)
        
        # Регистрируем шрифты для PDF
        self._register_fonts()
    
    def _register_fonts(self):
        """Регистрация шрифтов для PDF."""
        try:
            font_path = Path("resources/fonts/DejaVuSans.ttf")
            if font_path.exists():
                pdfmetrics.registerFont(TTFont('DejaVuSans', str(font_path)))
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', str(font_path)))
            else:
                logger.warning("Шрифт DejaVuSans не найден, используется стандартный")
        except Exception as e:
            logger.warning(f"Ошибка регистрации шрифтов: {e}")
    
    def export_to_pdf(self, request_id: int, template: str = "detailed") -> str:
        """
        Экспорт заявки в PDF.
        
        Args:
            request_id: ID заявки
            template: Шаблон экспорта ("detailed", "summary", "report")
            
        Returns:
            Путь к созданному файлу
        """
        try:
            # Получаем данные заявки
            request_data = self._get_request_data(request_id)
            if not request_data:
                raise ValidationError(f"Заявка {request_id} не найдена")
            
            # Создаем имя файла
            filename = f"request_{request_data['request_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.export_dir / filename
            
            # Создаем PDF документ
            doc = SimpleDocTemplate(str(file_path), pagesize=A4, leftMargin=20*mm, rightMargin=20*mm)
            story = []
            
            # Стили
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='DejaVuSans-Bold' if 'DejaVuSans-Bold' in pdfmetrics.getRegisteredFontNames() else 'Helvetica-Bold',
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=20
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName='DejaVuSans-Bold' if 'DejaVuSans-Bold' in pdfmetrics.getRegisteredFontNames() else 'Helvetica-Bold',
                fontSize=12,
                spaceBefore=10,
                spaceAfter=5
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='DejaVuSans' if 'DejaVuSans' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
                fontSize=10
            )
            
            # Заголовок
            story.append(Paragraph(f"Заявка на лабораторные испытания № {request_data['request_number']}", title_style))
            story.append(Spacer(1, 10*mm))
            
            # Информация о материале
            story.append(Paragraph("Информация о материале", heading_style))
            material_data = [
                ["Марка материала:", request_data['grade']],
                ["Номер плавки:", request_data['heat_num']],
                ["Размер:", request_data['size']],
                ["Тип проката:", request_data['rolling_type'] or "—"],
                ["Номер сертификата:", request_data['cert_num'] or "—"]
            ]
            
            material_table = Table(material_data, colWidths=[60*mm, 100*mm])
            material_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans' if 'DejaVuSans' in pdfmetrics.getRegisteredFontNames() else 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, black),
                ('BACKGROUND', (0, 0), (0, -1), (0.9, 0.9, 0.9))
            ]))
            
            story.append(material_table)
            story.append(Spacer(1, 10*mm))
            
            # Результаты испытаний
            if template == "detailed":
                story.append(Paragraph("Результаты испытаний", heading_style))
                
                for result in request_data['results']:
                    story.append(Paragraph(f"• {result['name']}", normal_style))
                    
                    # Форматируем результат
                    if isinstance(result['result'], dict):
                        for key, value in result['result'].items():
                            story.append(Paragraph(f"  {key}: {value}", normal_style))
                    else:
                        story.append(Paragraph(f"  Результат: {result['result']}", normal_style))
                    
                    story.append(Spacer(1, 5*mm))
            
            elif template == "summary":
                # Краткая сводка
                story.append(Paragraph("Краткая сводка", heading_style))
                
                summary_data = [["Испытание", "Результат"]]
                for result in request_data['results']:
                    if isinstance(result['result'], dict):
                        formatted_result = ", ".join([f"{k}: {v}" for k, v in result['result'].items()])
                    else:
                        formatted_result = str(result['result'])
                    summary_data.append([result['name'], formatted_result])
                
                summary_table = Table(summary_data, colWidths=[80*mm, 80*mm])
                summary_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans' if 'DejaVuSans' in pdfmetrics.getRegisteredFontNames() else 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, black),
                    ('BACKGROUND', (0, 0), (-1, 0), (0.8, 0.8, 0.8)),
                    ('FONTNAME', (0, 0), (-1, 0), 'DejaVuSans-Bold' if 'DejaVuSans-Bold' in pdfmetrics.getRegisteredFontNames() else 'Helvetica-Bold')
                ]))
                
                story.append(summary_table)
            
            # Информация о заявке
            story.append(Spacer(1, 10*mm))
            story.append(Paragraph("Информация о заявке", heading_style))
            
            request_info = [
                ["Дата создания:", request_data['created_at']],
                ["Статус:", request_data['status']],
                ["Сценарий:", request_data['scenario_name'] or "—"],
            ]
            
            request_table = Table(request_info, colWidths=[60*mm, 100*mm])
            request_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'DejaVuSans' if 'DejaVuSans' in pdfmetrics.getRegisteredFontNames() else 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, black),
                ('BACKGROUND', (0, 0), (0, -1), (0.9, 0.9, 0.9))
            ]))
            
            story.append(request_table)
            
            # Футер
            story.append(Spacer(1, 20*mm))
            story.append(Paragraph(f"Документ создан: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style))
            
            # Создаем PDF
            doc.build(story)
            
            logger.info(f"PDF создан: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Ошибка экспорта в PDF: {e}")
            raise BusinessLogicError(
                message="Ошибка экспорта в PDF",
                original_error=e
            )
    
    def export_to_docx(self, request_id: int, template: str = "detailed") -> str:
        """
        Экспорт заявки в DOCX.
        
        Args:
            request_id: ID заявки
            template: Шаблон экспорта
            
        Returns:
            Путь к созданному файлу
        """
        try:
            # Получаем данные заявки
            request_data = self._get_request_data(request_id)
            if not request_data:
                raise ValidationError(f"Заявка {request_id} не найдена")
            
            # Создаем имя файла
            filename = f"request_{request_data['request_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.export_dir / filename
            
            # Создаем документ
            doc = Document()
            
            # Заголовок
            title = doc.add_heading(f"Заявка на лабораторные испытания № {request_data['request_number']}", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Информация о материале
            doc.add_heading('Информация о материале', level=1)
            
            material_table = doc.add_table(rows=5, cols=2)
            material_table.style = 'Table Grid'
            
            material_data = [
                ("Марка материала:", request_data['grade']),
                ("Номер плавки:", request_data['heat_num']),
                ("Размер:", request_data['size']),
                ("Тип проката:", request_data['rolling_type'] or "—"),
                ("Номер сертификата:", request_data['cert_num'] or "—")
            ]
            
            for i, (label, value) in enumerate(material_data):
                material_table.cell(i, 0).text = label
                material_table.cell(i, 1).text = str(value)
                
                # Стилизация
                material_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
                material_table.cell(i, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Результаты испытаний
            doc.add_heading('Результаты испытаний', level=1)
            
            if template == "detailed":
                for result in request_data['results']:
                    doc.add_heading(result['name'], level=2)
                    
                    if isinstance(result['result'], dict):
                        results_table = doc.add_table(rows=len(result['result']), cols=2)
                        results_table.style = 'Table Grid'
                        
                        for i, (key, value) in enumerate(result['result'].items()):
                            results_table.cell(i, 0).text = key
                            results_table.cell(i, 1).text = str(value)
                    else:
                        p = doc.add_paragraph(f"Результат: {result['result']}")
            
            elif template == "summary":
                # Таблица сводки
                summary_table = doc.add_table(rows=len(request_data['results'])+1, cols=2)
                summary_table.style = 'Table Grid'
                
                # Заголовок
                summary_table.cell(0, 0).text = "Испытание"
                summary_table.cell(0, 1).text = "Результат"
                
                for i, (cell_0, cell_1) in enumerate([(summary_table.cell(0, 0), summary_table.cell(0, 1))]):
                    cell_0.paragraphs[0].runs[0].font.bold = True
                    cell_1.paragraphs[0].runs[0].font.bold = True
                
                # Данные
                for i, result in enumerate(request_data['results'], 1):
                    summary_table.cell(i, 0).text = result['name']
                    
                    if isinstance(result['result'], dict):
                        formatted_result = ", ".join([f"{k}: {v}" for k, v in result['result'].items()])
                    else:
                        formatted_result = str(result['result'])
                    
                    summary_table.cell(i, 1).text = formatted_result
            
            # Информация о заявке
            doc.add_heading('Информация о заявке', level=1)
            
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ("Дата создания:", request_data['created_at']),
                ("Статус:", request_data['status']),
                ("Сценарий:", request_data['scenario_name'] or "—")
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = str(value)
                info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            
            # Футер
            doc.add_paragraph()
            footer_p = doc.add_paragraph(f"Документ создан: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Сохраняем документ
            doc.save(str(file_path))
            
            logger.info(f"DOCX создан: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Ошибка экспорта в DOCX: {e}")
            raise BusinessLogicError(
                message="Ошибка экспорта в DOCX",
                original_error=e
            )
    
    def export_to_xlsx(self, request_ids: List[int], template: str = "data") -> str:
        """
        Экспорт заявок в XLSX.
        
        Args:
            request_ids: Список ID заявок
            template: Шаблон экспорта ("data", "analysis", "summary")
            
        Returns:
            Путь к созданному файлу
        """
        try:
            # Создаем имя файла
            filename = f"requests_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = self.export_dir / filename
            
            # Создаем рабочую книгу
            wb = Workbook()
            
            # Стили
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill("solid", fgColor="366092")
            border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin")
            )
            
            if template == "data":
                # Лист с данными заявок
                ws = wb.active
                ws.title = "Данные заявок"
                
                # Заголовки
                headers = [
                    "Номер заявки", "Дата создания", "Статус", "Марка материала",
                    "Номер плавки", "Размер", "Тип проката", "Сценарий"
                ]
                
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.border = border
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Данные
                for row, request_id in enumerate(request_ids, 2):
                    request_data = self._get_request_data(request_id)
                    if not request_data:
                        continue
                    
                    data = [
                        request_data['request_number'],
                        request_data['created_at'],
                        request_data['status'],
                        request_data['grade'],
                        request_data['heat_num'],
                        request_data['size'],
                        request_data['rolling_type'] or "—",
                        request_data['scenario_name'] or "—"
                    ]
                    
                    for col, value in enumerate(data, 1):
                        cell = ws.cell(row=row, column=col, value=value)
                        cell.border = border
                
                # Автоширина колонок
                for col in range(1, len(headers) + 1):
                    ws.column_dimensions[get_column_letter(col)].width = 15
                
                # Лист с результатами
                if len(request_ids) == 1:
                    request_data = self._get_request_data(request_ids[0])
                    if request_data and request_data['results']:
                        ws_results = wb.create_sheet("Результаты")
                        
                        # Заголовки
                        ws_results.cell(1, 1, "Испытание").font = header_font
                        ws_results.cell(1, 1).fill = header_fill
                        ws_results.cell(1, 2, "Результат").font = header_font
                        ws_results.cell(1, 2).fill = header_fill
                        
                        # Данные результатов
                        for row, result in enumerate(request_data['results'], 2):
                            ws_results.cell(row, 1, result['name'])
                            
                            if isinstance(result['result'], dict):
                                formatted_result = ", ".join([f"{k}: {v}" for k, v in result['result'].items()])
                            else:
                                formatted_result = str(result['result'])
                            
                            ws_results.cell(row, 2, formatted_result)
                        
                        # Стилизация
                        for row in range(1, len(request_data['results']) + 2):
                            for col in range(1, 3):
                                ws_results.cell(row, col).border = border
                        
                        ws_results.column_dimensions['A'].width = 30
                        ws_results.column_dimensions['B'].width = 50
            
            elif template == "analysis":
                # Лист для анализа данных
                ws = wb.active
                ws.title = "Анализ данных"
                
                # Собираем все результаты
                all_results = []
                for request_id in request_ids:
                    request_data = self._get_request_data(request_id)
                    if request_data:
                        for result in request_data['results']:
                            all_results.append({
                                'request_number': request_data['request_number'],
                                'grade': request_data['grade'],
                                'heat_num': request_data['heat_num'],
                                'test_name': result['name'],
                                'result': result['result']
                            })
                
                # Заголовки
                headers = ["Заявка", "Марка", "Плавка", "Испытание", "Результат"]
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.border = border
                
                # Данные
                for row, result in enumerate(all_results, 2):
                    formatted_result = result['result']
                    if isinstance(formatted_result, dict):
                        formatted_result = json.dumps(formatted_result, ensure_ascii=False)
                    
                    data = [
                        result['request_number'],
                        result['grade'],
                        result['heat_num'],
                        result['test_name'],
                        str(formatted_result)
                    ]
                    
                    for col, value in enumerate(data, 1):
                        cell = ws.cell(row=row, column=col, value=value)
                        cell.border = border
                
                # Автоширина колонок
                for col in range(1, len(headers) + 1):
                    ws.column_dimensions[get_column_letter(col)].width = 20
            
            # Сохраняем файл
            wb.save(str(file_path))
            
            logger.info(f"XLSX создан: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Ошибка экспорта в XLSX: {e}")
            raise BusinessLogicError(
                message="Ошибка экспорта в XLSX",
                original_error=e
            )
    
    def export_batch(self, request_ids: List[int], format: str = "pdf", 
                    template: str = "detailed") -> List[str]:
        """
        Пакетный экспорт множества заявок.
        
        Args:
            request_ids: Список ID заявок
            format: Формат экспорта ("pdf", "docx", "xlsx")
            template: Шаблон экспорта
            
        Returns:
            Список путей к созданным файлам
        """
        try:
            created_files = []
            
            if format == "xlsx":
                # Для XLSX создаем один файл со всеми заявками
                file_path = self.export_to_xlsx(request_ids, template)
                created_files.append(file_path)
            else:
                # Для PDF и DOCX создаем отдельные файлы
                for request_id in request_ids:
                    try:
                        if format == "pdf":
                            file_path = self.export_to_pdf(request_id, template)
                        elif format == "docx":
                            file_path = self.export_to_docx(request_id, template)
                        else:
                            raise ValidationError(f"Неподдерживаемый формат: {format}")
                        
                        created_files.append(file_path)
                        
                    except Exception as e:
                        logger.error(f"Ошибка экспорта заявки {request_id}: {e}")
                        continue
            
            logger.info(f"Пакетный экспорт завершен. Создано {len(created_files)} файлов")
            return created_files
            
        except Exception as e:
            logger.error(f"Ошибка пакетного экспорта: {e}")
            raise BusinessLogicError(
                message="Ошибка пакетного экспорта",
                original_error=e
            )
    
    def _get_request_data(self, request_id: int) -> Optional[Dict[str, Any]]:
        """Получение данных заявки для экспорта."""
        try:
            cursor = self.db_connection.cursor()
            cursor.execute("""
                SELECT lr.*, g.grade, rt.type AS rolling_type, m.size, m.heat_num, m.cert_num,
                       ts.name AS scenario_name
                FROM lab_requests lr
                JOIN Materials m ON lr.material_id = m.id
                JOIN Grades g ON m.grade_id = g.id
                LEFT JOIN RollingTypes rt ON m.rolling_type_id = rt.id
                LEFT JOIN test_scenarios ts ON lr.scenario_id = ts.id
                WHERE lr.id = ?
            """, (request_id,))
            
            row = cursor.fetchone()
            if not row:
                return None
            
            # Парсим JSON данные
            tests = json.loads(row['tests_json']) if row['tests_json'] else []
            results = json.loads(row['results_json']) if row['results_json'] else []
            
            return {
                'id': row['id'],
                'request_number': row['request_number'],
                'created_at': row['created_at'],
                'status': row['status'],
                'grade': row['grade'],
                'rolling_type': row['rolling_type'],
                'size': row['size'],
                'heat_num': row['heat_num'],
                'cert_num': row['cert_num'],
                'scenario_name': row['scenario_name'],
                'tests': tests,
                'results': results
            }
            
        except Exception as e:
            logger.error(f"Ошибка получения данных заявки {request_id}: {e}")
            return None
    
    def get_export_templates(self) -> Dict[str, List[str]]:
        """
        Получение списка доступных шаблонов экспорта.
        
        Returns:
            Словарь с шаблонами по форматам
        """
        return {
            'pdf': ['detailed', 'summary', 'report'],
            'docx': ['detailed', 'summary', 'official'],
            'xlsx': ['data', 'analysis', 'summary']
        }
    
    def cleanup_old_exports(self, days_old: int = 30) -> int:
        """
        Очистка старых экспортированных файлов.
        
        Args:
            days_old: Возраст файлов в днях для удаления
            
        Returns:
            Количество удаленных файлов
        """
        try:
            deleted_count = 0
            cutoff_time = datetime.now().timestamp() - (days_old * 24 * 60 * 60)
            
            for file_path in self.export_dir.glob("*"):
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        deleted_count += 1
                        logger.debug(f"Удален старый экспорт: {file_path}")
                    except OSError as e:
                        logger.warning(f"Не удалось удалить файл {file_path}: {e}")
            
            logger.info(f"Очистка экспортов завершена. Удалено {deleted_count} файлов")
            return deleted_count
            
        except Exception as e:
            logger.error(f"Ошибка очистки старых экспортов: {e}")
            return 0 